#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
book_parser.py —— 薄解析器：把 EPUB / 文字版 PDF 抽成 Qing 阅读分析法两段式所需的结构。

输出结构（与 Qing阅读分析法-两段式调用规格.md §1 一致）：
{
  "title":    str,
  "author":   str,
  "preface":  str,          # 序言 / 引言 / 推荐序 拼接（可选）
  "toc":      [             # 目录树，两级：章 → 小节
      {"level": 1, "title": "第1章 ...", "children": ["小节1", "小节2"]},
      ...
  ],
  "chapters": {             # 键 = 目录中 level=1 的标题，值 = 该章全文
      "第1章 ...": "……全文……",
      ...
  }
}

用法：
    python3 book_parser.py 书.epub            # 打印摘要
    python3 book_parser.py 书.epub -o out.json  # 另存 JSON

依赖：EPUB 仅用标准库；PDF 需 pip install pymupdf。
"""

import os
import re
import json
import html
import zipfile
import argparse
from html.parser import HTMLParser
from xml.etree import ElementTree as ET


# ---------------------------------------------------------------------------
# HTML -> 纯文本
# ---------------------------------------------------------------------------

class _TextExtractor(HTMLParser):
    """把 XHTML/HTML 抽成纯文本：跳过 script/style/head，块级标签换行。"""

    BLOCK_TAGS = {'p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                  'li', 'blockquote', 'tr', 'section', 'article'}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._skip = 0
        self._in_head = False
        self._parts = []

    def handle_starttag(self, tag, attrs):
        if tag in ('script', 'style'):
            self._skip += 1
        if tag == 'head':
            self._in_head = True
        if tag in self.BLOCK_TAGS:
            self._parts.append('\n')

    def handle_endtag(self, tag):
        if tag in ('script', 'style') and self._skip > 0:
            self._skip -= 1
        if tag == 'head':
            self._in_head = False
        if tag in self.BLOCK_TAGS:
            self._parts.append('\n')

    def handle_data(self, data):
        if self._skip == 0 and not self._in_head:
            self._parts.append(data)


def extract_html_text(raw):
    """HTML 字符串 -> 干净的纯文本。"""
    p = _TextExtractor()
    p.feed(raw)
    text = ''.join(p._parts)
    text = text.replace('\r\n', '\n').replace('\r', '\n')  # 统一换行符
    text = text.replace('\xa0', ' ')            # &nbsp;
    text = re.sub(r'[ \t\f\v]+', ' ', text)     # 压缩空白
    text = re.sub(r'\n[ \t]*\n+', '\n', text)   # 压空行
    return text.strip()


# ---------------------------------------------------------------------------
# 目录层级启发式（当 ncx 是扁平结构、没有嵌套 navPoint 时，用来区分"章"和"小节"）
# ---------------------------------------------------------------------------

# 一级章：带编号的章/课/讲/部分/篇，或英文 Chapter/Part
_RE_NUMBERED_CHAPTER = re.compile(
    r'^(第\s*[0-9一二三四五六七八九十百千]+\s*[章课讲部分篇]|'
    r'chapter\s*\d+|part\s*\d+|section\s*\d+)',
    re.IGNORECASE,
)

# 明确视为一级的前言/后记等
_TOP_LEVEL_KEYS = ('推荐序', '再版序', '自序', '代序', '新版序', '序言', '序', '前言', '引言', '导言',
                   '结语', '后记', '致谢', '附录', '参考文献')

# 纯结构页，不进目录、不进章节（封面/扉页/版权/目录页）
_SKIP_KEYS = ('封面', '扉页', '版权页', '版权', '目录')


def _is_numbered(title):
    """是否带编号的章/课/部分/篇（如「第1章」「Chapter 2」）——这是「书有章节结构」的信号。"""
    return bool(_RE_NUMBERED_CHAPTER.match(title))


def _is_front_back(title):
    """是否前言/后记类（序、前言、引言、结语、致谢、附录等）——这类在任何书里都独立成节。"""
    head = title[:4]
    return any(head.startswith(kw) for kw in _TOP_LEVEL_KEYS)


def _is_skip(title):
    head = title[:4]
    return any(head.startswith(kw) for kw in _SKIP_KEYS)


def _is_front_matter(title):
    head = title[:4]
    return any(head.startswith(kw) for kw in ('推荐序', '再版序', '自序', '代序', '新版序', '序言', '序', '前言', '引言', '导言'))


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------

def _local(tag):
    """去掉 XML 命名空间前缀，只留本地名。"""
    return tag.rsplit('}', 1)[-1] if '}' in tag else tag


def _findall(elem, name):
    """按本地名递归找所有子元素（命名空间无关）。"""
    return [e for e in elem.iter() if _local(e.tag) == name]


def _first(elem, name):
    r = _findall(elem, name)
    return r[0] if r else None


def _text_of(elem, name):
    e = _first(elem, name)
    return ''.join(e.itertext()).strip() if e is not None else ''


def _parse_opf(raw):
    """解析 OPF，返回 (metadata_dict, manifest{id:href}, spine[idref...], ncx_href, nav_href)。"""
    root = ET.fromstring(raw)
    meta = {
        'title': _text_of(root, 'title'),
        'author': _text_of(root, 'creator'),
        'publisher': _text_of(root, 'publisher'),
    }
    manifest = {}
    nav_href = None
    for item in _findall(root, 'item'):
        iid = item.get('id')
        href = item.get('href')
        if iid and href:
            manifest[iid] = href
            if item.get('properties') == 'nav' or 'nav' in (item.get('properties') or ''):
                nav_href = href

    spine = []
    spine_elem = _first(root, 'spine')
    if spine_elem is not None:
        ncx_id = spine_elem.get('toc')
        spine = [ir.get('idref') for ir in _findall(spine_elem, 'itemref') if ir.get('idref')]
    else:
        ncx_id = None

    ncx_href = manifest.get(ncx_id) if ncx_id else None
    if ncx_href is None:
        # 兜底：从 manifest 里按 media-type 找 ncx
        for item in _findall(root, 'item'):
            if item.get('media-type') == 'application/x-dtbncx+xml':
                ncx_href = item.get('href')
                break

    return meta, manifest, spine, ncx_href, nav_href


def _parse_ncx(raw):
    """解析 toc.ncx，返回有序的 [(title, src, depth)]。depth 由 navPoint 嵌套层级决定。"""
    root = ET.fromstring(raw)
    navmap = _first(root, 'navMap')
    if navmap is None:
        navmap = root
    out = []

    def walk(elems, depth):
        for np_ in elems:
            if _local(np_.tag) != 'navPoint':
                continue
            title = ''
            src = ''
            subs = []
            for child in np_:
                ln = _local(child.tag)
                if ln == 'navLabel':
                    t = _first(child, 'text')
                    title = ''.join(t.itertext()).strip() if t is not None else ''
                elif ln == 'content':
                    src = (child.get('src') or '').split('#')[0]
                elif ln == 'navPoint':
                    subs.append(child)
            if title:
                out.append((title, src, depth))
            walk(subs, depth + 1)

    walk(list(navmap), 0)
    return out


def parse_epub(path):
    zf = zipfile.ZipFile(path)

    # 1. 找 OPF 路径
    container = zf.read('META-INF/container.xml')
    opf_path = _first(ET.fromstring(container), 'rootfile').get('full-path')

    # 2. 解析 OPF
    base = os.path.dirname(opf_path)
    meta, manifest, spine, ncx_href, nav_href = _parse_opf(zf.read(opf_path))

    # 3. 读目录（优先 ncx，其次 nav.xhtml）
    entries = []  # [(title, src, depth)]
    if ncx_href:
        ncx_path = os.path.normpath(os.path.join(base, ncx_href)).replace(os.sep, '/')
        if ncx_path in zf.namelist():
            entries = [(t, s, d) for t, s, d in _parse_ncx(zf.read(ncx_path))]
    if not entries and nav_href:
        # EPUB3 nav 兜底：这里只做最简处理（取 <nav> 里的 <a href>）
        nav_path = os.path.normpath(os.path.join(base, nav_href)).replace(os.sep, '/')
        if nav_path in zf.namelist():
            entries = _parse_nav_xhtml(zf.read(nav_path))

    if not entries:
        # 完全没有目录：退回 spine 顺序，标题用文件名
        entries = [(manifest.get(ir, ir), manifest.get(ir, ir), 0) for ir in spine]

    # 4. 抽取各文件文本（缓存，src 可能被多章节共享/重复）
    text_cache = {}

    def get_text(src):
        if src in text_cache:
            return text_cache[src]
        full = os.path.normpath(os.path.join(base, src)).replace(os.sep, '/')
        if full in zf.namelist():
            text_cache[src] = extract_html_text(zf.read(full).decode('utf-8', 'ignore'))
        else:
            text_cache[src] = ''
        return text_cache[src]

    # 5. 组装：分组到一级章，拼正文
    # 先判断分组模式：
    #   nested        —— ncx 有嵌套（深度 ≥1），用深度分「章 / 小节」
    #   有编号章       —— 平铺但存在「第N章/Chapter N」等，用编号分（其余标题归为小节）
    #   无编号章       —— 散文集等，全部视为一级章（兜底）
    non_skip = [(t, s, d) for (t, s, d) in entries
                if t.strip() and not _is_skip(t.strip())]
    max_depth = max((d for _, _, d in non_skip), default=0)
    nested = max_depth >= 1
    has_numbered = any(_is_numbered(t) for t, _, _ in non_skip)
    flat_fallback = (not nested) and (not has_numbered) and bool(non_skip)

    toc = []          # [{"level":1,"title":..,"children":[..]}]
    chapters = {}     # 章标题 -> 全文
    preface = []
    cur = None        # 当前一级章索引
    for title, src, depth in entries:
        title = title.strip()
        if not title or _is_skip(title):
            continue
        if flat_fallback:
            is_top = True
        elif nested:
            is_top = (depth == 0)
        else:
            is_top = _is_numbered(title) or _is_front_back(title)

        if not is_top:
            # 视为小节，挂到当前章，正文并入
            if cur is not None:
                toc[cur]['children'].append(title)
                chapters[toc[cur]['title']] += '\n\n' + get_text(src)
            continue
        # 一级章
        cur = len(toc)
        toc.append({'level': 1, 'title': title, 'children': []})
        chapters[title] = get_text(src)
        if _is_front_matter(title):
            preface.append(get_text(src))

    # 合并 preface：序言正文里已含小节的话会重复，这里只取一级章自身文本
    preface_text = '\n\n'.join(p for p in preface if p).strip()

    return {
        'title': meta.get('title') or '',
        'author': meta.get('author') or '',
        'publisher': meta.get('publisher') or '',
        'preface': preface_text,
        'toc': toc,
        'chapters': chapters,
    }


def _parse_nav_xhtml(raw):
    """EPUB3 nav.xhtml 最简解析：提取 <nav> 内 <a href> 的文本与地址。"""
    from xml.etree import ElementTree as _ET
    try:
        root = _ET.fromstring(raw)
    except _ET.ParseError:
        root = _ET.fromstring('<html>' + raw + '</html>')
    out = []
    for a in root.iter():
        if _local(a.tag) == 'a' and a.text and a.get('href'):
            out.append((a.text.strip(), a.get('href').split('#')[0], 0))
    return out


# ---------------------------------------------------------------------------
# PDF（文字版，带书签）
# ---------------------------------------------------------------------------

def parse_pdf(path):
    """解析文字版、带书签的 PDF（扫描版需另做 OCR，不在本层处理）。用 PyPDF2（已普遍安装）。"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError(
            '解析 PDF 需要 PyPDF2（或 pymupdf），请先安装：pip install PyPDF2\n'
            '（注意：仅支持「文字版、带书签」的 PDF；扫描版 PDF 需另行 OCR，不在本层处理）'
        )

    reader = PdfReader(path)
    meta = reader.metadata or {}

    # 展平书签 → [(level, title, page_1indexed)]，level 0 = 章，≥1 = 小节
    entries = []

    def walk(items, level):
        for it in items:
            if isinstance(it, list):
                walk(it, level + 1)
            else:
                try:
                    pg = reader.get_destination_page_number(it) + 1
                except Exception:
                    continue  # 拿不到页码的条目跳过
                t = (it.title or '').strip()
                if t:
                    entries.append((level, t, pg))

    if reader.outline:
        walk(reader.outline, 0)

    total_pages = len(reader.pages)

    def page_text(p1):  # p1 为 1-indexed 页码
        i = p1 - 1
        return (reader.pages[i].extract_text() or '').strip() if 0 <= i < total_pages else ''

    # 一级条目（章 / 前言 / 后记），跳过纯结构页（封面/扉页/版权/目录）
    l0 = [(t, p) for (l, t, p) in entries if l == 0 and not _is_skip(t)]
    if not l0 and entries:
        # 书签无层级时兜底：把非 skip 的全部当一级
        l0 = [(t, p) for (l, t, p) in entries if not _is_skip(t)]

    toc = []
    chapters = {}
    preface = []
    for idx, (title, start) in enumerate(l0):
        end = l0[idx + 1][1] if idx + 1 < len(l0) else total_pages + 1
        # 章正文 = [start, end) 页，逐页拼接
        text = '\n'.join(page_text(p) for p in range(start, end) if page_text(p))
        children = [t for (l, t, p) in entries
                    if l >= 1 and not _is_skip(t) and start <= p < end]
        toc.append({'level': 1, 'title': title, 'children': children})
        chapters[title] = text
        if _is_front_matter(title):
            preface.append(text)

    title = meta.get('/Title') or meta.get('Title') or os.path.splitext(os.path.basename(path))[0]
    author = meta.get('/Author') or meta.get('Author') or ''
    return {
        'title': title,
        'author': author,
        'publisher': '',
        'preface': '\n\n'.join(p for p in preface if p).strip(),
        'toc': toc,
        'chapters': chapters,
    }


# ---------------------------------------------------------------------------
# 分发
# ---------------------------------------------------------------------------

def parse_txt(path):
    """纯文本：按「第N章 / Chapter N」等标题切章，无标题则整本作一章。"""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    re_ch = re.compile(
        r'^(第\s*[0-9一二三四五六七八九十百千]+\s*[章课讲部分篇])'
        r'|^(chapter\s+\d+|part\s+\d+|section\s+\d+)',
        re.IGNORECASE,
    )

    chapters = {}
    toc = []
    cur_title = None
    cur_parts = []
    front = []

    for line in text.split('\n'):
        s = line.strip()
        if re_ch.match(s):
            if cur_title is not None:
                chapters[cur_title] = '\n'.join(cur_parts)
            cur_title = s
            cur_parts = []
            toc.append({'level': 1, 'title': cur_title, 'children': []})
        elif cur_title is None:
            front.append(line)
        else:
            cur_parts.append(line)

    if cur_title is not None:
        chapters[cur_title] = '\n'.join(cur_parts)

    if not toc:
        body = text.strip()
        if body:
            toc = [{'level': 1, 'title': '正文', 'children': []}]
            chapters = {'正文': body}

    return {
        'title': os.path.splitext(os.path.basename(path))[0],
        'author': '',
        'publisher': '',
        'preface': '\n'.join(front).strip()[:3000],
        'toc': toc,
        'chapters': chapters,
    }


def _docx_heading_level(style):
    """把 Word 样式名映射为标题层级：0=正文，1=章，2=小节。"""
    s = (style or '').lower()
    if 'heading' in s or '标题' in s:
        m = re.search(r'(\d+)', s)
        return int(m.group(1)) if m else 1
    return 0


def parse_docx(path):
    """Word：按标题样式（Heading 1/标题 1 = 章，Heading 2/标题 2 = 小节）重建目录。"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError('解析 docx 需要 python-docx，请先安装：pip install python-docx')

    doc = Document(path)
    core = doc.core_properties

    chapters = {}
    toc = []
    cur_title = None
    cur_parts = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        text = para.text.strip()
        if not text:
            continue
        lvl = _docx_heading_level(style)
        if lvl == 1:
            if cur_title is not None:
                chapters[cur_title] = '\n'.join(cur_parts)
                cur_parts = []
            cur_title = text
            toc.append({'level': 1, 'title': text, 'children': []})
        elif lvl == 2 and cur_title is not None:
            toc[-1]['children'].append(text)
            cur_parts.append(text)
        else:
            cur_parts.append(text)

    if cur_title is not None:
        chapters[cur_title] = '\n'.join(cur_parts)

    if not toc:
        body = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        if body:
            toc = [{'level': 1, 'title': '正文', 'children': []}]
            chapters = {'正文': body}

    return {
        'title': (core.title or '').strip() or os.path.splitext(os.path.basename(path))[0],
        'author': (core.author or '').strip(),
        'publisher': '',
        'preface': '',
        'toc': toc,
        'chapters': chapters,
    }


def parse_mobi(path):
    """MOBI/AZW：优先用 calibre 的 ebook-convert 转成 EPUB 再解析；无 calibre 则报错。"""
    import subprocess
    epub_path = path + '.epub'
    try:
        subprocess.run(['ebook-convert', path, epub_path],
                       check=True, capture_output=True, timeout=180)
    except FileNotFoundError:
        raise RuntimeError('解析 mobi 需要 calibre（ebook-convert），服务器未安装，请先把 mobi 转成 EPUB')
    except Exception as e:
        raise RuntimeError(f'mobi 转 epub 失败：{e}')
    try:
        return parse_epub(epub_path)
    finally:
        if os.path.exists(epub_path):
            os.remove(epub_path)


def parse_book(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.epub':
        return parse_epub(path)
    if ext == '.pdf':
        return parse_pdf(path)
    if ext == '.txt':
        return parse_txt(path)
    if ext == '.docx':
        return parse_docx(path)
    if ext in ('.mobi', '.azw3', '.azw'):
        return parse_mobi(path)
    raise ValueError(f'暂不支持的格式：{ext}')


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(description='把书抽成 Qing 两段式所需的结构化数据')
    ap.add_argument('book', help='书文件路径（.epub 或文字版 .pdf）')
    ap.add_argument('-o', '--out', help='另存 JSON 的路径')
    args = ap.parse_args()

    data = parse_book(args.book)

    # 摘要
    print(f"书名：{data['title']}")
    print(f"作者：{data['author']}")
    print(f"序言：{len(data['preface'])} 字")
    print(f"章节数：{len(data['chapters'])}")
    print('-' * 50)
    for t in data['toc']:
        kids = f"（{len(t['children'])} 小节）" if t['children'] else ''
        n = len(data['chapters'].get(t['title'], ''))
        print(f"  {t['title']}  {n} 字 {kids}")

    if args.out:
        with open(args.out, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"\n已写入 {args.out}")


if __name__ == '__main__':
    main()
