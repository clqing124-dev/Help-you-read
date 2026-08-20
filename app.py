import os
import re
import uuid
from io import BytesIO
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS

import book_parser
import analyzer

app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB 限制

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 支持格式：EPUB / 文字版 PDF / TXT / DOCX / MOBI（mobi 依赖服务器装有 calibre）
ALLOWED_EXTENSIONS = {'.epub', '.pdf', '.txt', '.docx', '.mobi'}


@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze_book():
    """上传书籍 → 两段式（挑章节 + 出报告）分析。"""
    if not os.environ.get('DEEPSEEK_API_KEY'):
        return jsonify({'error': '服务器未配置 DEEPSEEK_API_KEY 环境变量，请联系管理员'}), 500

    if 'file' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({'error': f'暂不支持 {ext} 格式，请上传 EPUB、PDF、TXT、DOCX 或 MOBI'}), 400

    temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4().hex}{ext}")
    file.save(temp_path)

    try:
        report, sel, book = analyzer.analyze_book(temp_path, verbose=True)
        return jsonify({
            'success': True,
            'book_title': book.get('title') or os.path.splitext(file.filename)[0],
            'analysis': report,
            'selected_chapters': [ch.get('title', '') for ch in (sel.get('selected_chapters') or [])],
        })
    except Exception as e:
        return jsonify({'error': f'分析失败：{e}'}), 500
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def _clean_inline(text):
    """去掉行内标记（**加粗** `代码` *斜体*），保留纯文本。"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    return text.strip()


def _add_para(doc, text, style=None):
    """加段落，支持 **加粗** 行内标记。"""
    p = doc.add_paragraph(style=style)
    for part in re.split(r'(\*\*.+?\*\*)', text):
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part:
            p.add_run(part)
    return p


def _add_table(doc, rows):
    """把若干行「| a | b |」markdown 表格转成 docx 表格。"""
    data = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip('|').split('|')]
        if all(re.match(r'^:?-{2,}:?$', c) for c in cells):
            continue  # 表头分隔行
        data.append(cells)
    if not data:
        return
    ncols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(data):
        for ci in range(ncols):
            table.cell(ri, ci).text = _clean_inline(row[ci]) if ci < len(row) else ''


def markdown_to_docx(md):
    """把 Markdown 报告转成 docx Document。"""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            _add_table(doc, rows)
            continue
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            doc.add_heading(_clean_inline(m.group(2)), level=min(len(m.group(1)), 4))
            i += 1
            continue
        if re.match(r'^(-{3,}|\*{3,}|_{3,})$', s):
            i += 1
            continue
        if s.startswith('>'):
            p = _add_para(doc, _clean_inline(s.lstrip('>').strip()))
            p.paragraph_format.left_indent = Pt(24)
            i += 1
            continue
        if re.match(r'^[-*+]\s+', s):
            _add_para(doc, _clean_inline(re.sub(r'^[-*+]\s+', '', s)), style='List Bullet')
            i += 1
            continue
        m = re.match(r'^\d+[\.\)]\s+(.*)', s)
        if m:
            _add_para(doc, _clean_inline(m.group(1)), style='List Number')
            i += 1
            continue
        _add_para(doc, _clean_inline(s))
        i += 1
    return doc


@app.route('/export-docx', methods=['POST'])
def export_docx():
    """把 Markdown 报告转成 docx 文件下载。"""
    data = request.get_json(silent=True) or {}
    md = data.get('markdown', '')
    title = (data.get('title') or '阅读报告').strip()
    if not md:
        return jsonify({'error': '缺少报告内容'}), 400
    doc = markdown_to_docx(md)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf, as_attachment=True, download_name=f'{title}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
